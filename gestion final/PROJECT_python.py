from tkcalendar import Calendar
from tkinter.filedialog import *
from tkinter.ttk import *
from tkinter import ttk
import customtkinter as ctk
import re
from CTkMessagebox import CTkMessagebox
from mysql.connector import Error
from mysql.connector.errors import IntegrityError
from tkinter import *
from PIL import ImageTk, Image
import mysql.connector
import tkinter as tk
from docx import Document
from docx.shared import Inches
class Mainprogramme:
    global E1,E2,FRAME,CANVA,image
    def __init__(self,root,mycursor):
        #une image du font
        bg = ImageTk.PhotoImage(file="GESTIONDESETUDIANTS.png")
        self.CANVA = Canvas(root, width=7000, height=2500)
        self.CANVA.create_image(0, 0, image=bg,anchor='nw')
        self.CANVA.pack(fill=BOTH, expand=True)
        self.CANVA.bind("<Configure>", self.resize_image)
        #crée la premiere frame 
        self.FRAME = ctk.CTkFrame(master=self.CANVA,bg_color="black")
        self.FRAME.pack(pady=100, padx=200)
        self.FRAME.rowconfigure(0, weight=1)
        self.FRAME.columnconfigure(0, weight=1)
        self.FRAME.size()

        self.label1 = ctk.CTkLabel(master=self.FRAME, text="Login System", text_color="white", font=("Robot", 40))
        self.label1.pack(pady=60, padx=60)#(12,10)

        #######first image#################################
        image = Image.open("User.png")
        image = image.resize((15, 15))
        image = ImageTk.PhotoImage(image)
        LABELimage1= tk.Label(master=self.FRAME, image=image)
        # Keep a reference to the image to prevent it from being garbage collected
        LABELimage1.image = image
        # Place the label widget on the frame
        LABELimage1.place(x=85, y=172)
        #######second image#################################
        image = Image.open("Pass.png")
        image = image.resize((15, 15))
        image = ImageTk.PhotoImage(image)
        LABELimage = tk.Label(master=self.FRAME, image=image)
        # Keep a reference to the image to prevent it from being garbage collected
        LABELimage.image = image
        # Place the label widget on the frame
        LABELimage.place(x=85, y=212)
        ###################################################
        self.E1=ctk.CTkEntry(master=self.FRAME, placeholder_text="Username")
        self.E1.pack(pady=0, padx=10)
        self.E=ctk.CTkEntry(master=self.FRAME, placeholder_text="Password", show="*")
        self.E.pack(pady=10, padx=10)
        self.button = ctk.CTkButton(master=self.FRAME, text="login", command=lambda:self.login(mycursor,root))
        self.button.pack(pady=5, padx=10)
        self.button = ctk.CTkButton(master=self.FRAME, text="Sign in ", command=lambda:self.Sign_in(mycursor,root))
        self.button.pack(pady=5, padx=10)
        self.button = ctk.CTkButton(master=self.FRAME, text="Suprimmer un étudiant",command=lambda: self.Supprimer(mycursor,root))
        self.button.pack(pady=5, padx=1)
        root.mainloop()




    def resize_image(self, event):
        global image, image2
        # Resize the image with the new width and height
        image = Image.open("GESTIONDESETUDIANTS.png")
        resized = image.resize((event.width, event.height))
        # Update the image on the canvas
        self.CANVA.delete("all")  # Remove previous image
        image2 = ImageTk.PhotoImage(resized)
        self.CANVA.create_image(0, 0, anchor="nw", image=image2)






    def login(self,mycursor,root):#le button login <------------------------
        name = self.E1.get()
        password = self.E.get()
        query =  f"select * from ETUDIANT where Nom=\"{name}\" and Passsword=\"{password}\""
        mycursor.execute(query)
        print("true")
        # fetcher le resultat
        result = mycursor.fetchall()
        if len(result) == 0:
            CTkMessagebox(title="Warning", message="Username ou Password incorrect ")

        else:
            self.FRAME.destroy()
            self.CANVA.bind("<Configure>", self.resize_image)
            self.frame2 = ctk.CTkScrollableFrame(master=self.CANVA, width=1200, height=600,scrollbar_button_color="purple", scrollbar_button_hover_color="deeppink")
            self.frame2.pack(padx=10,pady=10)
            image = Image.open("User.png")
            image = image.resize((40, 40))
            image = ImageTk.PhotoImage(image)
            sLABELIMAGE=tk.Label(master=self.frame2,image=image)
            sLABELIMAGE.image = image
            label2 = ctk.CTkLabel(master=self.frame2,text="Bienvenue sur la plateforme de Gestion des Etudiants!\n vous pouvez consultez vos informations:)",text_color="white", font=("Robot", 15))
            label2.grid(row=3, column=2, padx=10, pady=10)
            self.BUTTON = ctk.CTkButton(master=self.frame2, text="Vos Informations ",command=lambda: self.Information(mycursor, name, password))
            self.BUTTON.grid(row=4, column=1, padx=10, pady=10)
            self.BUTTON = ctk.CTkButton(master=self.frame2, text="Ajouter votre photo ", command=self.ajouter)
            self.BUTTON.grid_rowconfigure(0, weight=1)
            self.BUTTON.grid_columnconfigure(0, weight=1)
            self.BUTTON.grid(row=4, column=3, padx=10, pady=10)
            #une image pour le profil
            image = Image.open("User.png")
            image = image.resize((100, 100))
            image = ImageTk.PhotoImage(image)
            Label1 = tk.Label(master=self.frame2, image=image)
            Label1.grid_rowconfigure(0, weight=1)
            Label1.grid_columnconfigure(0, weight=1)
            Label1.image = image
            Label1.grid(row=5, column=3, padx=1, pady=1)
            Label2 = tk.Label(master=self.frame2, text="Vous pouvez consulter quelque cours dans votre domaine",bg="black",fg="white")
            Label2 .grid(row=7, column=1,padx=2,pady=2)
            self.BUTTON= ctk.CTkButton(master=self.frame2, text="Télécharger Vos informations",command=lambda: self.Imprimer(mycursor, name, password))
            self.BUTTON.grid(row=7, column=3,padx=10,pady=10)
            self.BUTTON = ctk.CTkButton(master=self.frame2, text="Le cour du java",command=self.Telecharger)
            self.BUTTON.grid(row=8, column=1, padx=1, pady=5)
            self.BUTTON = ctk.CTkButton(master=self.frame2, text=" Le cour du python",command=self.Telecharger1)
            self.BUTTON.grid(row=9, column=1, padx=1, pady=5)
            self.BUTTON = ctk.CTkButton(master=self.frame2, text="Le cour de langage SQL",command=self.Telecharger2)
            self.BUTTON.grid(row=10, column=1, padx=1, pady=5)
            self.BUTTON = ctk.CTkButton(master=self.frame2, text="Le cour de PL",command=self.Telecharger3)
            self.BUTTON.grid(row=11, column=1, padx=1, pady=1)
            self.BUTTON = ctk.CTkButton(master=self.frame2, text="Quitter", command=lambda:self.EXIT2(root))
            self.BUTTON.grid(row=8, column=3,padx=1,pady=1)






    def Information(self,mycursor,name,password):#le button vos information<----------------------
        #selecter les informations à partir de base de données
        query=f"select * from ETUDIANT where Nom=\"{name}\" and Passsword=\"{password}\" "
        mycursor.execute(query)
        information = mycursor.fetchall()
        #row=[cne,cin,nom,prenom,numero,email,password,nv,datedenaissance]
        for row in information:
          CTl1 = ctk.CTkLabel(master=self.frame2,text=f"LE NOM:   \"{row[2]}\"\nLE PRENOM :   \"{row[3]}\"\nCIN:    \"{row[1]}\"\nCNE :   \"{row[0]}\"\nLE NUMERO: \"{row[4]}\"\nL'EMAIL :   \"{row[5]}\"\nLE MOT DE PASSE DE Compte  :   \"{row[6]}\"\nLE NIVEAU SCOLAIRE:   \"{row[7]}\"\n DATE DE NAISSANCE:  \"{row[8]}\"\n", text_color="white",bg_color="purple",font=("Robot",20))
          CTl1.grid(row=5,column=1,padx=1,pady=1)






        
    def ajouter(self):#le button "ajouter une photode profil"<---------------------
        filepath= askopenfilename(title="Ouvrir une Image", filetypes=[('png files', '.png'), ('all files', '.*')])
        image = Image.open(filepath)
        image = image.resize((100, 100))
        image = ImageTk.PhotoImage(image)
        Label1 = tk.Label(master=self.frame2, image=image)
        Label1.grid_rowconfigure(0, weight=1)
        Label1.grid_columnconfigure(0, weight=1)
        Label1.image = image
        Label1.grid(row=5, column=3, padx=1, pady=1)







    def Telecharger(self):#le button" lecour java"<---------------------
        import webbrowser
        webbrowser.open('https://www.emse.fr/~picard/cours/1A/java/livretJava.pdf', new=2)
    
    
    
    
    
    def Telecharger1(self):#le button" lecour python"<---------------------
        import webbrowser
        webbrowser.open('https://cfm.ehu.es/ricardo/docs/python/Learning_Python.pdf', new=2)
    
    
    def Telecharger2(self):#le button" lecourSQL"<---------------------
        import webbrowser
        webbrowser.open('https://sql.sh/ressources/cours-sql-sh-.pdf', new=2)



    def Telecharger3(self):#le button" lecour PLSQL"<---------------------
        import webbrowser
        webbrowser.open('https://www.lri.fr/~fiorenzi/Teaching/BDAS/C2.pdf', new=2)



        
    def Imprimer(self,mycursor,name,password):#le button" Télécharger vos information"<---------------------
        filepath = askopenfilename(title="Ouvrir une Image", filetypes=[('png files', '.png'), ('all files', '.*')])
        query = f"select * from ETUDIANT where Nom=\"{name}\" and Passsword=\"{password}\" "
        mycursor.execute(query)
        information = mycursor.fetchall()
        doc=Document()
        import os
        for row in information:
            p=doc.add_paragraph()
            r=p.add_run()
            r.add_text(f"LE NOM:   \"{row[2]}\"\nLE PRENOM :   \"{row[3]}\"\nCIN:    \"{row[1]}\"\nCNE :   \"{row[0]}\"\nLE NUMERO: \"{row[4]}\"\nL'EMAIL :   \"{row[5]}\"\nLE MOT DE PASSE DE Compte  :   \"{row[6]}\"\nLE NIVEAU SCOLAIRE:   \"{row[7]}\"\n DATE DE NAISSANCE:  \"{row[8]}\"\n")
            r.add_picture(filepath,width=Inches(3))
            doc.save('vos_information.docx')
            CTkMessagebox(title="Validé", message="vous pouvez consultez votre fichier \n dans le méme dossier ")
            


    
    
    
    def Sign_in(self,mycursor,root):#le button" Sign in"<---------------------
        self.FRAME.destroy()
        self.CANVA.bind("<Configure>", self.resize_image)
        self.FRAME = ctk.CTkFrame(master=self.CANVA, bg_color="black", width=1350, height=700)
        self.FRAME.pack(padx=10,pady=10)
        self.FRAME.grid_columnconfigure(0, weight=1)
        labell = ctk.CTkLabel(master=self.FRAME, text="Obtenez votre compte gratuit dès maintenant", text_color="white",font=("Robot", 20))
        labell.grid(row=1, column=0,padx=20)
        label_text = ctk.CTkLabel(self.FRAME, text="Veuillez remplir tous les champs obligatoires*", text_color="deeppink",font=("String var",13))
        label_text.grid(row=2, column=0,padx=70)
        label = ctk.CTkLabel(self.FRAME, text="Nom *")#<----------------------------
        label.grid(row=6, column=0)
        self.entry1 = ctk.CTkEntry(master=self.FRAME)
        self.entry1.grid(row=7, column=0)
        label1 = ctk.CTkLabel(self.FRAME, text="Prenom *")
        label1.grid(row=9, column=0)
        self.entry2 = ctk.CTkEntry(master=self.FRAME)
        self.entry2.grid(row=10, column=0)
        label2 = ctk.CTkLabel(self.FRAME, text="Email *")
        label2.grid(row=12, column=0)
        self.entry3 = ctk.CTkEntry(master=self.FRAME)
        self.entry3.grid(row=13, column=0)
        label3 = ctk.CTkLabel(self.FRAME, text="CIN *")
        label3.grid(row=15, column=0)
        self.entry4 = ctk.CTkEntry(master=self.FRAME)
        self.entry4.grid(row=16, column=0)
        label4 = ctk.CTkLabel(self.FRAME, text="CNE *")
        label4.grid(row=18, column=0)
        self.entry5 = ctk.CTkEntry(master=self.FRAME)
        self.entry5.grid(row=19, column=0)
        label5 = ctk.CTkLabel(self.FRAME, text="Niveau Scolaire *" )
        label5.grid(row=6, column=1)
        self.les_choix=ttk.Combobox(self.FRAME,font=("Arial", 10))
        self.les_choix['values']=['CP1','CP2','ID1','ID2','ID3']
        self.les_choix.grid(row=7, column=1)
        self.les_choix.configure(background='black', foreground='white')
        label6 = ctk.CTkLabel(self.FRAME, text="Mot de passe *")
        label6.grid(row=9, column=1)
        self.entry7 = ctk.CTkEntry(master=self.FRAME,show='*')
        self.entry7.grid(row=10, column=1)
        label7 = ctk.CTkLabel(self.FRAME, text="le mot de passe confirmè *")
        label7.grid(row=12, column=1)
        self.entry8 = ctk.CTkEntry(master=self.FRAME, show='*')
        self.entry8.grid(row=13, column=1)
        label8 = ctk.CTkLabel(self.FRAME, text="Numero *")
        label8.grid(row=15, column=1)
        self.entry9 = ctk.CTkEntry(master=self.FRAME)
        self.entry9.grid(row=16, column=1)
        label10 = ctk.CTkLabel(self.FRAME, text="Date de Naissance *")
        label10.grid(row=18, column=1)
        self.entry10 = ctk.CTkEntry(master=self.FRAME)
        self.entry10.grid(row=19, column=1)
        self.entry10.insert(0, "yyyy-mm-dd")
        self.entry10.bind("<1>", self.pick_date)
        self.button = ctk.CTkButton(master=self.FRAME, text="Soumettre et mémoriser", command=lambda:self.insert(mycursor,root))
        self.button.grid(padx=20, pady=20, row=36, column=0)
        self.button = ctk.CTkButton(master=self.FRAME, text="Retourner", command=lambda:self.Return(mycursor,root))
        self.button.grid(padx=30, pady=20, row=36, column=1)



    def pick_date(self, event):#entry 10<---------------------
        global cal, date_window
        self.date_window = Toplevel()
        self.date_window .wait_visibility()  
        self.date_window.grab_set()
        self.date_window.title('chosir une date')
        self.date_window.geometry('250x220+590+370')
        cal =Calendar(self.date_window,selectmode="day", date_pattern="yyyy-mm-dd")
        cal.place(x=0, y=0)
        submit_btn = Button(self.date_window, text="Soumettre", command=self.grab_date)
        submit_btn.place(x=80, y=190)

    def grab_date(self):
        self.entry10.delete(0, END)
        self.entry10.insert(0, cal.get_date())
        self.date_window.destroy()

    
    
    
    
    def insert(self,mycursor,root):#le button" Soumettre et mémpriser"<---------------------
        #cette fonction vas vérifier l'input de l'utilisateur ,et aprés vas l'insérer à la base de données
        #ces dictionnaires sont déclarés pour optimiser le code 
        D={}
        B={}
        nom = self.entry1.get()
        prenom = self.entry2.get()
        Email = self.entry3.get()
        CIN = self.entry4.get()
        CNE = self.entry5.get()
        Numero = self.entry9.get()
        password = self.entry7.get()
        nv = self.les_choix.get()
        cp = self.entry8.get()
        datenaissance = self.entry10.get()
        #pour email
        regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b'
        #le dictionnaire D contient la position exacte de la ligne d'instruction d'erreur qui sera nécessaire ci-dessous
        D[nom] = [8, 0]
        D[prenom] = [11, 0]
        D[Email] = [14, 0]
        D[CIN] = [17, 0]
        D[CNE] = [20, 0]
        D[nv] = [8, 1]
        D[password] = [11, 1]
        # le dictionnaire B contient les variables ENTRY qui seront nécessaires ci-dessous
        B[nom] = self.entry1
        B[prenom] = self.entry2
        B[Email] = self.entry3
        B[CIN] = self.entry4
        B[CNE] = self.entry5
        B[nv] = self.les_choix
        B[password] = self.entry7
        # Verfication0: vérifier si les valeurs dans D sont numeric digits
        for i in D:
            if  i.isdigit() == True:
                  B[i].delete(0, END)
                  lab = ctk.CTkLabel(self.FRAME, text="ne doit pas contenir seulement des chiffres", text_color="red")
                  lab.grid(row=D[i][0], column=D[i][1])
        #verfication1: si le champ est vide.
        if nom == "" or prenom == "" or Email == "" or CIN == "" or Numero == "" or CIN == "" or password == "" or cp == "" or nv == "":
            label_erreur = ctk.CTkLabel(self.FRAME, text="Erreur : Veuillez remplir tous les champs obligatoires*",fg_color="red")
            label_erreur.grid(row=3, column=0)
        # verfication2: si le numero est caractere
        #verfication2: si le numero est caractere
        elif Numero.isdigit() == False:
            self.entry9.delete(0, END)
            lab = ctk.CTkLabel(self.FRAME, text="ne doit pas contenir seulement des caractères", text_color="red")
            lab.grid(row=17, column=1)
        #verification3 de mot de pass confirme
        elif cp!=password:
            self.entry8.delete(0, END)
            lab = ctk.CTkLabel(self.FRAME, text="vous devez taper le même Mot de passe", text_color="red")
            lab.grid(row=14, column=1)
        elif not (re.fullmatch(regex, Email)):
            self.entry8.delete(0, END)
            lab = ctk.CTkLabel(self.FRAME, text="email invalid", text_color="red")
            lab.grid(row=14, column=0)
        else:
            # vérifier si il a déja un compte
            query1 = f"select * from ETUDIANT where  CNE=\"{CNE}\""
            mycursor.execute(query1)
            print("true")
            # fetching the result
            result1 = mycursor.fetchall()
            if len(result1)!=0 :
                CTkMessagebox(title="Warning ", message="vous avez deja un compte")
            else:
                # Exécuter le code pour insérer des informations dans la base de données
                query = f"insert into ETUDIANT values (\"{CNE}\",\"{CIN}\",\"{nom}\",\"{prenom}\",\"{Numero}\",\"{Email}\",\"{password}\",\"{nv}\",\"{datenaissance}\")"
                mycursor.execute(query)
                mycursor.execute("commit")
                CTkMessagebox(title="Warning", message="l'opération est bien réussie")
                # créer une autre page pour  log in ou exit
                ctk.set_appearance_mode("dark")
                ctk.set_default_color_theme("green")
                Root = ctk.CTk()
                Root.title('login')
                Root.geometry("300x300")

                Root.configure(background='purple')
                frame1 = ctk.CTkFrame(master=Root, width=250, height=250)
                frame1.place(x=25, y=30)
                L = ctk.CTkLabel(master=frame1, text="Acceder votre compte", text_color="purple", font=("Robot", 15))
                L.place(x=45, y=40)
                b = ctk.CTkButton(master=frame1, text="Connecter ", command=lambda:self.Remember(mycursor,Root,frame1))
                b.place(x=50, y=80)
                b = ctk.CTkButton(master=frame1, text="Quiter", command=lambda:self.EXIT1(Root,root))
                b.place(x=50, y=130)
                Root.mainloop()

    
    
    
    
    def Return(self,mycursor,root):#le button "Rotourner" dans le Sign_in<----------------
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("green")
        Root = ctk.CTk()
        Root.geometry("400x300")
        Root.title('Gestion des etudiants')
        Root.configure(background='purple')
        frame = ctk.CTkFrame(master=Root, bg_color="purple")
        frame.pack(fill="both", expand=True)
        label1 = ctk.CTkLabel(master=frame, text="Login System", text_color="purple", font=("Robot", 42))
        label1.pack(pady=12, padx=10)
        self.E1= ctk.CTkEntry(master=frame, placeholder_text="Username")
        self.E1.pack(pady=12, padx=10)
        self.E= ctk.CTkEntry(master=frame, placeholder_text="Password", show="*")
        self.E.pack(pady=10, padx=10)
        button = ctk.CTkButton(master=frame, text="Connecter", command=lambda:self.login(mycursor,root))
        button.pack(pady=12, padx=10)
        Root.mainloop()
        Root.destroy()
        

    
    
    
    def Remember(self,mycursor,root,Frame):#le button "Soumaitre et mémoriser" dans le Sign_in<-------------
        Frame.destroy()
        frame = ctk.CTkFrame(master=root, bg_color="purple")
        frame.pack(fill="both", expand=True)
        label1 = ctk.CTkLabel(master=frame, text="Login System", text_color="white", font=("Robot", 30))
        label1.pack(pady=12, padx=10)
        self.E1 = ctk.CTkEntry(master=frame, placeholder_text="Username")
        self.E1.pack(pady=12, padx=10)
        self.E = ctk.CTkEntry(master=frame, placeholder_text="Password", show="*")
        self.E.pack(pady=10, padx=10)
        self.button = ctk.CTkButton(master=frame, text="login", command=lambda:self.login(mycursor,root))
        self.button.pack(pady=12, padx=10)
        nom=self.entry1.get()
        PAssword= self.entry7.get()
        self.E1.insert(0,nom)
        self.E.insert(0,PAssword)
        root.mainloop()
        root.destroy()



    def EXIT2(self, root):
        root.destroy()

        
    def EXIT1(self,Root,root):
        root.destroy()
        Root.destroy()




    def Supprimer(self, mycursor,root):#le button "Supprimer un étudiant" dans le login<----------------- 
        self.FRAME.destroy()
        self.FRAME= ctk.CTkFrame(master=self.CANVA, bg_color="black")
        self.FRAME.pack(pady=150, padx=200)
        self.FRAME.size()
        label = ctk.CTkLabel(master=self.FRAME, text="Entrez le CNE de l\'étudiant \n que voue voulez supprimer",font=("Robot", 20))
        label.pack(pady=50, padx=60)
        entry = ctk.CTkEntry(master=self.FRAME, placeholder_text="CNE")
        entry.pack(pady=15)
        button = ctk.CTkButton(master=self.FRAME, text="Supprimer", command=lambda: self.supp(mycursor, entry,root))
        button.pack(pady=15)
        button = ctk.CTkButton(master=self.FRAME, text="Rotourner",command=lambda:self.Return(mycursor,root))
        button.pack(pady=15)






    def supp(self,mycursor,entry, root):
        cne = entry.get()
        try:
            query1 = f"SELECT nom FROM ETUDIANT WHERE cne='{cne}'"
            mycursor.execute(query1)
            result = mycursor.fetchone()
            if not result:
                # If the student record does not exist, display an error message
                CTkMessagebox(title ="Erreur", message=f"L'étudiant avec le CNE '{cne}' n'existe pas.")
                return

            # Delete the student record
            query2 = f"DELETE FROM ETUDIANT WHERE cne='{cne}'"
            mycursor.execute(query2)
            mycursor.execute("COMMIT;")

            # Display a success message
            CTkMessagebox(title ="Opération réussie",message= "L'opération est validée.")

        except IntegrityError:
            # If the delete operation fails due to a foreign key constraint violation,
            # display an error message indicating the references to the student in other tables
            CTkMessagebox(title ="Erreur",message="Impossible de supprimer l'étudiant. Veuillez supprimer les références à cet étudiant dans les autres tables avant de réessayer.")

        except Error:
            # Catch other MySQL errors and display an error message
           CTkMessagebox(title ="Erreur",message="Une erreur s'est produite lors de la suppression de l'étudiant. Veuillez réessayer plus tard.")
        
        






        
######le programme######<----------------------------------------------------------------
#une into à l'application
ppage1=ctk.CTk()
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("green")
ppage1.geometry("600X300")#600X300
ppage1.title('Processus de L\'instalation')
frame = ctk.CTkFrame(master=ppage1)
frame.pack(pady=15, padx=20, fill="both",expand=True)
frame.size()
label= ctk.CTkLabel(master=frame, text="Bienvenue! vous étes dans l'Apk de Gestion des Etudiants\n\n Pour utiliser ce programme  vous devez installer les logiciels: \n MySQL et un IDE de PYTHON (Pycharm, VScode...) \n et les modules suivants :\n customtkinter,tkcalendar,CTkMessagebox,tkinter,PIL,mysql.connector,python-docx \n pour installer les modules utiliser votre console et la commande suivante:\n (pip install[nom du module])",font=("Robot",13),text_color="White",width=100,height=150)
label.pack(pady=10, padx=10)
button = ctk.CTkButton(master=frame, text="Suivant", command=lambda: suivant(frame,ppage1))
button.pack(pady=10, padx=10)


def suivant(event1,event2):
    event1.destroy()
    frame1=ctk.CTkFrame(master=event2, width=100,height=150)
    frame1.pack(pady=10, padx=20, fill="both",expand=True)
    frame1.size()
    label1= ctk.CTkLabel(frame1, text="Nom d\'utilisateur (MySql)")
    label1.grid(padx=150,row=16, column=0)
    entry1 = ctk.CTkEntry(master=frame1,placeholder_text="ex:root")
    entry1.grid(padx=150,row=17, column=0)
    

    label2 = ctk.CTkLabel(frame1, text="Nom de base de données")
    label2.grid(padx=150,row=18, column=0)
    entry2 = ctk.CTkEntry(master=frame1,placeholder_text="ex:sys")
    entry2.grid(padx=150,row=19, column=0)
    

    label3 = ctk.CTkLabel(frame1, text="Mot de passe")
    
    label3.grid(padx=150,row=20, column=0)
    entry3= ctk.CTkEntry(master=frame1,show="*",placeholder_text="ex:xxxxx")
    entry3.grid(padx=150,row=21, column=0)
    
    button = ctk.CTkButton(master=frame1, text="Suivant", command=lambda: connection(entry1,entry2,entry3,event2))
    button.grid(padx=150,pady=10,row=27, column=0)






#laconnection à la base de données
def connection(entry1, entry2, entry3, event2):
    ###get information####
    NU = entry1.get()
    BD = entry2.get()
    MP = entry3.get()

    try:

        mydb = mysql.connector.connect(
            host="localhost",
            user=NU,
            password=MP,
            database=BD
        )
        if mydb.is_connected():
            print("connection succecfully")
            mycursor = mydb.cursor()
            ######creation du table#######
            mycursor.execute(
                "create  table if not exists ETUDIANT(CNE varchar(10) primary key,CIN varchar(20),Nom varchar(20) ,Prenom varchar(20),Numero integer(15),Email varchar(40),Passsword varchar (30),NIVEAU_SCOLAIRE varchar(100),DATE_NAISSANCE DATE)")
            mycursor.execute("commit")
            event2.destroy()
            ctk.set_appearance_mode("dark")
            ctk.set_default_color_theme("green")
            root = ctk.CTk()
            root.geometry("800x650")  # 700x500
            root.title('Gestion des Etudiants')
            obj = Mainprogramme(root, mycursor)


    except mysql.connector.Error as e:
        print("false")
        CTkMessagebox(title="Warning", message="l'opération est tombée, veuillez vérifier vos informations")


ppage1.mainloop()


   


    

    
    

       
   





